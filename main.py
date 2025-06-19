from transformers import BartForConditionalGeneration, BartTokenizer, AutoModelForSequenceClassification, AutoTokenizer
from scipy.special import softmax
import torch
import requests
from docx import Document
from datetime import datetime

# (BART)
bart_model_name = "facebook/bart-large-cnn"
bart_tokenizer = BartTokenizer.from_pretrained(bart_model_name)
bart_model = BartForConditionalGeneration.from_pretrained(bart_model_name)

#  FinBERT
sentiment_model_name = "yiyanghkust/finbert-tone"
sentiment_tokenizer = AutoTokenizer.from_pretrained(sentiment_model_name)
sentiment_model = AutoModelForSequenceClassification.from_pretrained(sentiment_model_name)

#  Şirketler ve CEO'lar
companies = {
    "Apple": "Tim Cook",
    "Tesla": "Elon Musk",
    "Google": "Sundar Pichai",
    "Amazon": "Andy Jassy",
    "Microsoft": "Satya Nadella"
}

def summarize(text, max_length=60):
    if not text.strip():
        return ""

    inputs = bart_tokenizer(
        text,
        return_tensors="pt",
        truncation=True,
        max_length=1024,
        padding='longest'
    )

    summary_ids = bart_model.generate(
        inputs["input_ids"],
        num_beams=4,
        max_length=max_length,
        early_stopping=True
    )

    return bart_tokenizer.decode(summary_ids[0], skip_special_tokens=True)

def analyze_sentiment(text):
    if not text.strip():
        return 0.0
    try:
        inputs = sentiment_tokenizer(text, return_tensors="pt", truncation=True)
        with torch.no_grad():
            outputs = sentiment_model(**inputs)
        logits = outputs.logits[0].numpy()
        probs = softmax(logits)

        # FinBERT label order: [positive, neutral, negative]
        score = probs[0] - probs[2]  # pozitif - negatif
        return round(float(score), 2)
    except Exception as e:
        print("⚠️ Sentiment analizi hatası:", e)
        return 0.0

def extract_related_companies(text):
    text = text.lower()
    matched = set()
    for name, ceo in companies.items():
        if name.lower() in text or ceo.lower() in text:
            matched.add(name)
    return matched

def fetch_news(query, max_pages=1):
    NEWSDATA_API_KEY = "NEWSDATA_API_KEY"
    NEWSDATA_ENDPOINT = "https://newsdata.io/api/1/news"
    articles, next_page = [], None
    base_params = {
        "apikey": NEWSDATA_API_KEY,
        "q": query,
        "language": "en",
        "category": "business",
    }

    for _ in range(max_pages):
        params = base_params.copy()
        if next_page:
            params["page"] = next_page

        response = requests.get(NEWSDATA_ENDPOINT, params=params)
        if response.status_code != 200:
            print("API Hatası:", response.text)
            break

        data = response.json()
        results = data.get("results", [])
        if not results:
            break

        articles.extend(results)
        next_page = data.get("nextPage")
        if not next_page:
            break

    return articles

def save_report(scores, news_list, total_articles, top_n=5):
    # pozitif ve negatif filtrele
    sorted_news_pos = sorted(
        [n for n in news_list if n['sentiment'] > 0],
        key=lambda x: x['sentiment'],
        reverse=True
    )[:top_n]

    sorted_news_neg = sorted(
        [n for n in news_list if n['sentiment'] < 0],
        key=lambda x: x['sentiment']
    )[:top_n]


    doc = Document()
    doc.add_heading('Stock Analyse AI - Haftalık Rapor', 0)
    doc.add_paragraph(f"Rapor tarihi: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Toplam haber sayısı: {total_articles}")

    doc.add_heading('Şirket Bazlı Toplam Duygu Skorları', level=1)
    for c, s in scores.items():
        doc.add_paragraph(f"{c}: {s:+.2f}")

    doc.add_heading(f'En Pozitif {top_n} Haber', level=1)
    for news in sorted_news_pos:
        doc.add_paragraph(f"{news['company']} - {news['title']} ({news['sentiment']:+.2f})")
        doc.add_paragraph(news['url'])
        doc.add_paragraph('')

    doc.add_heading(f'En Negatif {top_n} Haber', level=1)
    for news in sorted_news_neg:
        doc.add_paragraph(f"{news['company']} - {news['title']} ({news['sentiment']:+.2f})")
        doc.add_paragraph(news['url'])
        doc.add_paragraph('')

    dosya_adi = f"Stock_Analyse_Rapor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(dosya_adi)
    print(f"\n✅ Rapor dosyası oluşturuldu: {dosya_adi}")

def fetch_and_score():
    scores = {name: 0.0 for name in companies}
    important_news = []
    seen_urls = set()
    total_articles = 0

    # Etkiyi artırabilecek anahtar kelimeler
    impact_keywords = ["regulation", "lawsuit", "fine", "settlement", "investigation", "acquisition", "merger", "antitrust", "earnings", "SEC", "compliance"]

    for company, ceo in companies.items():
        query = f"{company} OR {ceo}"
        articles = fetch_news(query, max_pages=2)

        for art in articles:
            url = art.get("link", "").strip()
            if not url or url in seen_urls:
                continue
            seen_urls.add(url)
            total_articles += 1

            title = (art.get("title") or "").strip()
            desc = (art.get("description") or "").strip()
            content = f"{title}. {desc}"

            related = extract_related_companies(content)
            if company not in related:
                continue

            summary = summarize(content)
            base_score = analyze_sentiment(summary if summary else content)

            #  Etki skoru ağırlıklandırma
            multiplier = 1.0
            title_lower = title.lower()

            if company.lower() in title_lower:
                multiplier *= 1.2
            if ceo.lower() in title_lower:
                multiplier *= 1.1
            if any(keyword in title_lower for keyword in impact_keywords):
                multiplier *= 1.5

            final_score = round(base_score * multiplier, 2)

            if abs(final_score) < 0.1:
                continue  # önemsiz

            print(f"\n📰 {title}")
            print(f"🔗 {url}")
            print(f"📝 Özet: {summary}")
            print(f"🎯 Sentiment: {base_score:+.2f} → Etki Skoru: {final_score:+.2f} (Çarpan: x{multiplier:.2f}) — {company}")

            scores[company] += final_score
            important_news.append({
                "company": company,
                "title": title,
                "url": url,
                "sentiment": final_score,
                "content": summary
            })

    print(f"\n📊 Toplam işlenen haber: {total_articles}")
    print("📊 Şirket Bazlı Toplam Skorlar:")
    for c, s in scores.items():
        print(f"{c}: {s:+.2f}")

    save_report(scores, important_news, total_articles)


if __name__ == "__main__":
    fetch_and_score()
